"""DOCX file parser using python-docx (MIT license).

Extracts text from paragraphs, tables, **and** embedded images (via OCR).
Images are OCR'd at their exact position in the document so that
context is preserved (e.g. a diagram between two paragraphs stays
inline rather than being appended at the end).
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

from rag_kb.parsers.base import DocumentParser, ParsedDocument

logger = logging.getLogger(__name__)


class DocxParser:
    """Parse Microsoft Word .docx files."""

    supported_extensions: list[str] = [".docx"]

    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("python-docx is required for DOCX parsing: pip install python-docx") from exc

        doc = Document(str(file_path))
        parts: list[str] = []
        images_ocrd = 0

        # Walk the document body children in order — this includes
        # both paragraphs (<w:p>) and tables (<w:tbl>), preserving
        # the original reading order.
        from docx.oxml.ns import qn

        for element in doc.element.body:
            tag = element.tag

            # --- Paragraph ---
            if tag == qn("w:p"):
                para_text, img_count = self._process_paragraph(element, doc)
                images_ocrd += img_count
                if para_text.strip():
                    parts.append(para_text.strip())

            # --- Table ---
            elif tag == qn("w:tbl"):
                from docx.table import Table
                table = Table(element, doc)
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)

        text = "\n\n".join(parts)

        metadata: dict[str, str] = {
            "format": "docx",
            "paragraph_count": str(len(doc.paragraphs)),
        }
        if images_ocrd:
            metadata["images_ocrd"] = str(images_ocrd)

        return ParsedDocument(
            text=text,
            source_path=str(file_path),
            metadata=metadata,
        )

    # ------------------------------------------------------------------

    @staticmethod
    def _process_paragraph(para_element, doc) -> tuple[str, int]:
        """Extract text from a paragraph, OCR-ing any inline images.

        Returns ``(combined_text, image_count)`` where *combined_text*
        interleaves normal text runs with ``[Image OCR]`` blocks so
        that the reading order is preserved.
        """
        from docx.oxml.ns import qn

        segments: list[str] = []
        img_count = 0

        for child in para_element:
            tag = child.tag

            # Regular text run
            if tag == qn("w:r"):
                # Check if this run contains a drawing (inline image)
                drawings = child.findall(f".//{qn('w:drawing')}")
                if drawings:
                    for drawing in drawings:
                        img_text, success = DocxParser._ocr_drawing(drawing, doc)
                        if success:
                            img_count += 1
                        if img_text:
                            segments.append(f"[Image OCR]\n{img_text}")
                else:
                    # Normal text
                    text_el = child.find(qn("w:t"))
                    if text_el is not None and text_el.text:
                        segments.append(text_el.text)

        return " ".join(segments) if segments else "", img_count

    @staticmethod
    def _ocr_drawing(drawing_element, doc) -> tuple[str, bool]:
        """OCR an image referenced by a ``<w:drawing>`` element.

        Resolves the relationship ID from the ``<a:blip>`` element,
        fetches the image blob from the document part, and runs OCR.

        Returns ``(ocr_text, success_bool)``.
        """
        from docx.oxml.ns import qn
        from PIL import Image as PILImage
        from rag_kb.parsers.image_parser import ocr_image

        try:
            blip = drawing_element.find(f".//{qn('a:blip')}")
            if blip is None:
                return "", False

            # r:embed holds the relationship ID for the image
            ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            r_embed = blip.get(f"{{{ns_r}}}embed")
            if not r_embed:
                return "", False

            related_part = doc.part.related_parts.get(r_embed)
            if related_part is None:
                return "", False

            pil_img = PILImage.open(io.BytesIO(related_part.blob)).convert("RGB")

            # Skip tiny images (icons, bullets)
            if pil_img.width < 50 or pil_img.height < 50:
                return "", False

            ocr_text, engine = ocr_image(pil_img, label=f"docx-image/{r_embed}")
            return ocr_text, bool(ocr_text)

        except Exception as exc:
            logger.debug("DOCX image OCR failed: %s", exc)
            return "", False
